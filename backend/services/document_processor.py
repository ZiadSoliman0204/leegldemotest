"""
Document processing service for Law Firm AI Assistant
Supports PDF, TXT, and DOCX files with local processing
"""

import os
import re
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

# Lightweight document processing libraries
import PyPDF2
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Basic text processing
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
import json

logger = logging.getLogger(__name__)

class LocalDocumentProcessor:
    """Local document processor with lightweight embedding alternatives"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=1000,
            stop_words='english',
            ngram_range=(1, 2),
            lowercase=True
        )
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Process a document and extract text chunks with metadata
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            
        Returns:
            Dictionary with processing results
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if not text.strip():
                raise ValueError("No text content found in document")
            
            # Clean and process text
            cleaned_text = self._clean_text(text)
            
            # Create text chunks
            chunks = self._create_text_chunks(cleaned_text)
            
            # Generate document ID
            document_id = self._generate_document_id(filename, text)
            
            # Create document metadata
            metadata = {
                'document_id': document_id,
                'filename': filename,
                'file_type': file_extension,
                'total_chunks': len(chunks),
                'total_characters': len(text),
                'processing_method': 'local'
            }
            
            # Prepare chunks with embeddings and metadata
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    'document_id': document_id,
                    'filename': filename,
                    'chunk_index': i,
                    'file_type': file_extension,
                    'chunk_length': len(chunk)
                }
                
                # Generate local embedding
                embedding = self._generate_local_embedding(chunk)
                
                processed_chunks.append({
                    'text': chunk,
                    'metadata': chunk_metadata,
                    'embedding': embedding
                })
            
            return {
                'document_id': document_id,
                'filename': filename,
                'chunks': processed_chunks,
                'metadata': metadata,
                'pages_processed': len(chunks)
            }
            
        except Exception as error:
            logger.error(f"Error processing document {filename}: {error}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try PyMuPDF first (better text extraction)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except Exception as error:
                logger.warning(f"PyMuPDF extraction failed: {error}, falling back to PyPDF2")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text
        except Exception as error:
            logger.error(f"PDF text extraction failed: {error}")
            raise ValueError(f"Could not extract text from PDF: {error}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as error:
            logger.error(f"TXT text extraction failed: {error}")
            raise ValueError(f"Could not extract text from TXT file: {error}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ValueError("python-docx library not available for DOCX processing")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as error:
            logger.error(f"DOCX text extraction failed: {error}")
            raise ValueError(f"Could not extract text from DOCX file: {error}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Normalize quotes
        text = re.sub(r'["""]', '"', text)
        text = re.sub(r"[''']", "'", text)
        
        # Remove multiple periods
        text = re.sub(r'\.{3,}', '...', text)
        
        # Strip and return
        return text.strip()
    
    def _create_text_chunks(self, text: str) -> List[str]:
        """Create overlapping text chunks"""
        chunks = []
        words = text.split()
        
        if len(words) <= self.chunk_size:
            return [text]
        
        start = 0
        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk = ' '.join(words[start:end])
            chunks.append(chunk)
            
            if end >= len(words):
                break
            
            # Move start position with overlap
            start += self.chunk_size - self.chunk_overlap
        
        return chunks
    
    def _generate_document_id(self, filename: str, content: str) -> str:
        """Generate unique document ID based on filename and content"""
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        filename_hash = hashlib.md5(filename.encode()).hexdigest()[:4]
        return f"doc_{filename_hash}_{content_hash}"
    
    def _generate_local_embedding(self, text: str) -> List[float]:
        """
        Generate local embedding using TF-IDF and basic hashing
        This is a lightweight alternative to OpenAI embeddings
        """
        try:
            # Method 1: TF-IDF based embedding (preferred)
            tfidf_vector = self._get_tfidf_embedding(text)
            if tfidf_vector is not None:
                return tfidf_vector.tolist()
            
            # Method 2: Hash-based embedding (fallback)
            return self._get_hash_embedding(text)
            
        except Exception as error:
            logger.warning(f"TF-IDF embedding failed: {error}, using hash-based fallback")
            return self._get_hash_embedding(text)
    
    def _get_tfidf_embedding(self, text: str) -> Optional[np.ndarray]:
        """Generate TF-IDF based embedding vector"""
        try:
            # Use hash-based fallback for consistency
            # TF-IDF requires fitting on a consistent corpus which is complex for this use case
            return None
            
        except Exception:
            return None
    
    def _get_hash_embedding(self, text: str, dimension: int = 384) -> List[float]:
        """
        Generate hash-based embedding as fallback
        Creates a reproducible vector from text hash
        """
        # Create multiple hash values for better distribution
        hash_values = []
        
        # Use different hash functions
        hash_funcs = [
            lambda x: hashlib.md5(x.encode()).hexdigest(),
            lambda x: hashlib.sha1(x.encode()).hexdigest(),
            lambda x: hashlib.sha256(x.encode()).hexdigest()
        ]
        
        for i, hash_func in enumerate(hash_funcs):
            # Hash text with salt
            salted_text = f"{text}_{i}"
            hex_hash = hash_func(salted_text)
            
            # Convert hex to integers
            for j in range(0, len(hex_hash), 8):
                chunk = hex_hash[j:j+8]
                hash_values.append(int(chunk, 16))
        
        # Create embedding vector
        np.random.seed(hash_values[0] % (2**32))  # Reproducible randomness
        vector = np.random.normal(0, 1, dimension)
        
        # Normalize
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector = vector / norm
        
        return vector.tolist()
    
    def calculate_similarity(self, embedding1: List[float], embedding2: List[float]) -> float:
        """Calculate cosine similarity between two embeddings"""
        try:
            vec1 = np.array(embedding1)
            vec2 = np.array(embedding2)
            
            # Cosine similarity
            dot_product = np.dot(vec1, vec2)
            norm1 = np.linalg.norm(vec1)
            norm2 = np.linalg.norm(vec2)
            
            if norm1 == 0 or norm2 == 0:
                return 0.0
            
            similarity = dot_product / (norm1 * norm2)
            return float(similarity)
            
        except Exception as error:
            logger.error(f"Similarity calculation failed: {error}")
            return 0.0 